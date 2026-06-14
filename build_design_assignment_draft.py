from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "design_assignment_draft.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="DADCE0", size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    values = {"top": top, "start": start, "bottom": bottom, "end": end}
    for m, v in values.items():
        tag = "w:{}".format(m)
        element = margins.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            margins.append(element)
        element.set(qn("w:w"), str(v))
        element.set(qn("w:type"), "dxa")


def style_paragraph(paragraph, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    style_paragraph(p, after=3)
    run = p.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    style_paragraph(p, after=16)
    run = p.add_run(subtitle)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    style_paragraph(p, before=18 if level == 1 else 12, after=6)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0 if level < 3 else 67)
    run.font.size = Pt(20 if level == 1 else 16 if level == 2 else 14)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        style_paragraph(p, after=4)
        run = p.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(11)


def add_placeholder(doc, title, lines=3):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16.5)
    cell = table.cell(0, 0)
    cell.width = Cm(16.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(cell)
    set_cell_margins(cell)
    set_cell_shading(cell, "F8F9FA")
    p = cell.paragraphs[0]
    style_paragraph(p, after=4)
    run = p.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)
    for _ in range(lines):
        p = cell.add_paragraph()
        style_paragraph(p, after=0)
        r = p.add_run(" ")
        r.font.size = Pt(12)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    add_title(
        doc,
        "Этюды: карта сфер",
        "Черновой PDF-каркас научно-проектного лонгрида для задания по дисциплине «Дизайн»",
    )

    add_heading(doc, "Большой блок с идеей", 1)
    idea = [
        "Один из участников проекта ранее провел визуальное исследование интернет-эстетик. По итогам этого исследования возникла необходимость переопределить и уточнить, чем являются интернет-эстетики и как они проявляют себя в цифровой среде.",
        "В ходе исследования были проанализированы более 200 стилей, представленных на Aesthetics Wiki, а также подробно рассмотрены более 40 эстетик. Это позволило сделать несколько выводов. Во-первых, при конструировании эстетик используются повторяющиеся визуальные приемы, задача которых состоит не только в оформлении изображения, но и в создании ауры, атмосферы или эмоционального режима восприятия. Во-вторых, четко разграничить принадлежность эстетик к определенному жанру почти невозможно: одни и те же визуальные признаки могут переходить между стилями, а часть контента в социальных сетях устроена по принципу микроэстетик, даже если такие эстетики не имеют отдельного названия.",
        "Из этого следует, что микроэстетики не стоит понимать только как институционализированные субкультурные направления с устойчивыми границами. Интернет-эстетика может быть рассмотрена шире: как изображение или мультимодальное явление, направленное на передачу особой атмосферы. Названная эстетика может состоять из большой подборки изображений, но потенциально может быть выражена и одной картинкой. Поэтому для данного проекта важнее не продолжать каталогизацию эстетик и не предлагать пользователю жестко выбирать между готовыми идентичностями вроде cottagecore или Barbiecore, а создать пространство, где минимальной единицей контента становится сфера.",
        "Сфера в приложении «Этюды» понимается как способ зафиксировать сиюминутное настроение, атмосферу воспоминания, пережитое событие или фантазию. Она может собираться из текста, изображения, музыки, места, времени, людей, предметов, времени дня, цветов и чувств. Такой формат позволяет перенести исследование интернет-эстетик из режима классификации в режим личной и коллективной атмосферной картографии.",
        "В мае 2026 года наша команда студентов DH-центра подготовила предложение для культурной институции о проведении выставки, посвященной интернет-эстетикам 2015-2017 годов. В качестве цифрового расширения выставки была предложена идея QR-перехода на специальную доску в формате лендинга, где посетители могли бы добавлять фотографии и тексты с собственными воспоминаниями и представлениями об эпохе. В приложении «Этюды» мы развиваем схожий потенциал: пользователь может не только просматривать готовую карту атмосфер, но и создавать собственные сферы как элементы личного архива, памяти и воображения.",
        "В рамках данного проекта мы фокусируемся не на всем приложении, а на одном его фрагменте: разделе «Карта». Именно здесь сфера становится видимой единицей интерфейса. На карте сферы существуют как миниатюры, расположенные в условной невидимой сетке с небольшими смещениями, создающими ощущение живого, не до конца упорядоченного поля. При нажатии сфера плавно увеличивается и раскрывается как прямоугольная область со скругленными углами: сверху может располагаться изображение, которое через мягкий градиент переходит к текстовым полям. Нижняя часть сферы предполагает прокрутку, позволяющую рассматривать подпись, ощущения, предметы, дату, время, место, запахи и музыкальную ассоциацию. Нажатие на пустые области по краям возвращает пользователя к карте, а горизонтальный свайп в раскрытом режиме позволяет переходить к соседним сферам без промежуточного возврата.",
    ]
    for paragraph in idea:
        add_body(doc, paragraph)

    add_placeholder(doc, "Место для короткой схемы: исследование → проблема классификации → сфера → карта", 4)

    add_heading(doc, "1. Введение", 1)
    add_body(doc, "Здесь будет краткая постановка проектной задачи: как визуальное исследование интернет-эстетик переводится в дизайн-концепцию интерфейса.")
    add_placeholder(doc, "Пустое место: 1-2 абзаца после доработки", 5)

    add_heading(doc, "2. Исследовательское основание", 1)
    add_body(doc, "Здесь нужно показать, какое визуальное исследование стало материалом проекта: корпус эстетик, принципы анализа, ключевые выводы.")
    add_placeholder(doc, "Пустое место: примеры эстетик, ссылки, изображения из исследования", 8)

    add_heading(doc, "3. Проектная идея: сфера вместо каталога", 1)
    add_body(doc, "Здесь раскрывается переход от готовых категорий эстетик к сфере как единице атмосферной фиксации.")
    add_placeholder(doc, "Пустое место: уточнение определения сферы и примеры возможных сфер", 6)

    add_heading(doc, "4. Аудитория и сценарий применения", 1)
    add_body(doc, "PDF готовится как научно-проектный лонгрид для защиты. Внутри текста также нужно обозначить внешнюю аудиторию самого проекта: пользователи «Этюдов», посетители культурного события, DH-сообщество или смешанная аудитория.")
    add_placeholder(doc, "Пустое место: финальное описание аудитории", 6)

    add_heading(doc, "5. Роли участников", 1)
    add_body(doc, "Проект предполагает групповую работу 3-4 человек. Роли нужно закрепить за конкретными участниками и обосновать.")
    add_bullets(doc, [
        "Исследователь: основание проекта, визуальное исследование, работа с источниками.",
        "Дизайнер: интерфейсная форма, карта, раскрытие сферы, прототип.",
        "Критик: культурная и медиакритическая рамка проекта.",
        "Прототипирование: работающий макет фрагмента приложения.",
    ])
    add_placeholder(doc, "Пустое место: имена участников и индивидуальные обоснования ролей", 7)

    add_heading(doc, "6. Интерфейсная концепция", 1)
    add_body(doc, "Раздел «Карта» показывает миниатюры сфер. Карта строится на невидимых прямоугольных областях: сферы меньше этих областей примерно в два раза и слегка смещаются, чтобы сохранить ощущение органической неупорядоченности.")
    add_body(doc, "При нажатии сфера плавно увеличивается, превращаясь в карточку со скругленными углами. Сверху может находиться изображение, ниже -- текстовые поля. Свайп влево или вправо в раскрытом режиме переключает соседние сферы.")
    add_placeholder(doc, "Место для изображения: скетч карты сфер", 9)
    add_placeholder(doc, "Место для изображения: раскрытая сфера", 9)
    add_placeholder(doc, "Место для изображения: экран создания сферы", 9)

    add_heading(doc, "7. Создание сферы", 1)
    add_body(doc, "В первой версии экран создания сферы содержит добавление фото и поля: подпись, ощущение, предметы, дата, время, место, запахи, название песни.")
    add_placeholder(doc, "Пустое место: решение, какие поля обязательные, а какие можно пропустить", 5)

    add_heading(doc, "8. Первый прототип", 1)
    add_body(doc, "Здесь будет описан работающий макет: что уже реализовано, что является заглушкой, какие решения проверяются на прототипе.")
    add_placeholder(doc, "Пустое место: скриншоты и комментарии к первому прототипу", 10)

    add_heading(doc, "9. Ценность результата", 1)
    add_body(doc, "Здесь нужно объяснить, как проект может применяться выбранной аудиторией и какую ценность он имеет: личный архив, карта атмосфер, цифровое расширение выставки, способ говорить о микроэстетиках без жесткой классификации.")
    add_placeholder(doc, "Пустое место: итоговые выводы", 6)

    add_heading(doc, "10. Использование ИИ", 1)
    add_body(doc, "В финальной версии здесь необходимо указать использованные промпты и кратко описать, как ИИ применялся в подготовке текста, структуры и/или макета.")
    add_placeholder(doc, "Пустое место: список промптов", 8)

    doc.save(OUT)


if __name__ == "__main__":
    build()
